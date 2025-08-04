from flask import Flask, request, send_file
from docx import Document
import difflib
import uuid
import os

app = Flask(__name__)

@app.route('/')
def home():
    return "✅ NDA Redline API is running."

@app.route('/compare-nda', methods=['POST'])
def compare_nda():
    file1 = request.files['template']
    file2 = request.files['counterparty']

    def get_text(file):
        doc = Document(file)
        return [p.text for p in doc.paragraphs if p.text.strip()]

    t_text = get_text(file1)
    c_text = get_text(file2)

    diff = difflib.HtmlDiff().make_file(
        t_text, c_text,
        fromdesc='Template NDA',
        todesc='Counterparty NDA'
    )

    redlined = Document()
    redlined.add_paragraph("Redlined NDA Comparison")
    for line in diff.splitlines():
        if line.strip():
            redlined.add_paragraph(line)

    # 🔒 Save file to /tmp so Render allows it
    filename = f"redlined_{uuid.uuid4().hex}.docx"
    filepath = os.path.join("/tmp", filename)
    redlined.save(filepath)

    # ✅ Set correct MIME type and force download name
    return send_file(
        filepath,
        as_attachment=True,
        download_name="Redlined NDA.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
